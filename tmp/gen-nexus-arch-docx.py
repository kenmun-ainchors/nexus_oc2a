#!/usr/bin/env python3
"""
Generate Nexus-System-Architecture-v1.0.docx from the Markdown source.
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_PATH = "/Users/ainchorsangiefpl/.openclaw/workspace/docs/Nexus-System-Architecture-v1.0.md"
DOCX_PATH = "/Users/ainchorsangiefpl/.openclaw/workspace/docs/Nexus-System-Architecture-v1.0.docx"

# Colours
BLUE = RGBColor(9, 105, 218)      # #0969da
DARK = RGBColor(30, 30, 30)
GREY = RGBColor(100, 100, 100)
TABLE_HEADER_BG = "0969da"
TABLE_ALT_BG    = "f0f4f8"


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_styled_paragraph(doc, text, style_name, bold=False, color=None, size=None):
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return p


def parse_inline(run_text):
    """Return list of (text, bold, code) tuples from inline markdown."""
    parts = []
    # Handle **bold** and `code`
    pattern = re.compile(r'(\*\*(.+?)\*\*|`(.+?)`)')
    last = 0
    for m in pattern.finditer(run_text):
        if m.start() > last:
            parts.append((run_text[last:m.start()], False, False))
        if m.group(0).startswith('**'):
            parts.append((m.group(2), True, False))
        else:
            parts.append((m.group(3), False, True))
        last = m.end()
    if last < len(run_text):
        parts.append((run_text[last:], False, False))
    return parts


def add_inline_paragraph(doc, text, style_name='Normal'):
    p = doc.add_paragraph(style=style_name)
    # Strip leading/trailing whitespace
    text = text.strip()
    # Remove any markdown link syntax [text](url) -> text
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    parts = parse_inline(text)
    for part_text, bold, code in parts:
        run = p.add_run(part_text)
        run.bold = bold
        if code:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(180, 60, 60)
    return p


def process_table(doc, header_row, data_rows):
    """Add a Word table from header_row (list) and data_rows (list of lists)."""
    cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=cols)
    table.style = 'Table Grid'

    # Header
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header_row):
        hdr_cells[i].text = h.strip()
        set_cell_bg(hdr_cells[i], TABLE_HEADER_BG)
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Data
    for ri, row in enumerate(data_rows):
        row_cells = table.rows[ri + 1].cells
        for ci, cell_text in enumerate(row):
            cell_text = cell_text.strip()
            # Remove markdown link
            cell_text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', cell_text)
            para = row_cells[ci].paragraphs[0]
            parts = parse_inline(cell_text)
            for part_text, bold, code in parts:
                run = para.add_run(part_text)
                run.bold = bold
                run.font.size = Pt(9)
                if code:
                    run.font.name = 'Courier New'
                    run.font.color.rgb = RGBColor(180, 60, 60)
            if ri % 2 == 1:
                set_cell_bg(row_cells[ci], TABLE_ALT_BG)

    doc.add_paragraph()  # spacing after table


def build_docx(md_path, docx_path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Font defaults
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_lines = []
    table_buffer = []
    in_table = False

    def flush_table():
        nonlocal table_buffer, in_table
        if not table_buffer:
            return
        # Parse table rows
        rows = []
        for tl in table_buffer:
            cells = [c for c in tl.strip().split('|') if True]
            # Remove leading/trailing empty from split
            if cells and cells[0].strip() == '':
                cells = cells[1:]
            if cells and cells[-1].strip() == '':
                cells = cells[:-1]
            rows.append(cells)
        # Filter separator rows (--- only)
        filtered = [r for r in rows if not all(re.match(r'^[-: ]+$', c.strip()) for c in r)]
        if len(filtered) >= 2:
            process_table(doc, filtered[0], filtered[1:])
        elif len(filtered) == 1:
            process_table(doc, filtered[0], [])
        table_buffer = []
        in_table = False

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Code block
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(style='Normal')
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(50, 50, 50)
                p.paragraph_format.left_indent = Inches(0.3)
                code_lines = []
                in_code_block = False
            else:
                if in_table:
                    flush_table()
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table detection
        stripped = line.strip()
        if stripped.startswith('|') and stripped.endswith('|'):
            in_table = True
            table_buffer.append(line)
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        # Blank line
        if not stripped:
            i += 1
            continue

        # Blockquote
        if stripped.startswith('>'):
            content = stripped.lstrip('> ').strip()
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.left_indent = Inches(0.3)
            parts = parse_inline(content)
            for pt, bold, code in parts:
                run = p.add_run(pt)
                run.bold = bold
                run.font.color.rgb = GREY
                run.font.size = Pt(9)
                if code:
                    run.font.name = 'Courier New'
            i += 1
            continue

        # Headings
        if stripped.startswith('#'):
            m = re.match(r'^(#{1,4})\s+(.*)', stripped)
            if m:
                level = len(m.group(1))
                heading_text = m.group(2).strip()
                # Remove anchor tags like {#anchor}
                heading_text = re.sub(r'\{#[^}]+\}', '', heading_text).strip()

                if level == 1:
                    p = doc.add_heading(level=0)
                    p.clear()
                    run = p.add_run(heading_text)
                    run.font.size = Pt(20)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(20, 20, 80)
                elif level == 2:
                    p = doc.add_heading(level=1)
                    p.clear()
                    run = p.add_run(heading_text)
                    run.font.size = Pt(15)
                    run.font.bold = True
                    run.font.color.rgb = BLUE
                elif level == 3:
                    p = doc.add_heading(level=2)
                    p.clear()
                    run = p.add_run(heading_text)
                    run.font.size = Pt(12)
                    run.font.bold = True
                    run.font.color.rgb = DARK
                else:
                    p = doc.add_heading(level=3)
                    p.clear()
                    run = p.add_run(heading_text)
                    run.font.size = Pt(11)
                    run.font.bold = True
                    run.font.color.rgb = DARK
                i += 1
                continue

        # Horizontal rule
        if stripped in ('---', '***', '___') or re.match(r'^-{3,}$', stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Bullet list
        if stripped.startswith('- ') or stripped.startswith('* '):
            content = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            parts = parse_inline(content)
            for pt, bold, code in parts:
                run = p.add_run(pt)
                run.bold = bold
                run.font.size = Pt(10)
                if code:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(180, 60, 60)
            i += 1
            continue

        # Numbered list
        if re.match(r'^\d+\.\s', stripped):
            content = re.sub(r'^\d+\.\s', '', stripped)
            p = doc.add_paragraph(style='List Number')
            parts = parse_inline(content)
            for pt, bold, code in parts:
                run = p.add_run(pt)
                run.bold = bold
                run.font.size = Pt(10)
                if code:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(180, 60, 60)
            i += 1
            continue

        # Normal paragraph
        add_inline_paragraph(doc, stripped)
        i += 1

    # Flush any remaining table
    if in_table:
        flush_table()

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == '__main__':
    build_docx(MD_PATH, DOCX_PATH)
