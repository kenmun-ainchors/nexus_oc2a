#!/usr/bin/env python3
"""Generate BASE1 Restore Runbook DOCX for Aevlith Nexus Platform."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ──────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x24, 0x29, 0x2F)

BLUE = RGBColor(0x09, 0x69, 0xDA)
GRAY = RGBColor(0x57, 0x60, 0x6A)

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_code_block(doc, text):
    for line in text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BLUE
    return h

# ═══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Aevlith Nexus Platform')
run.font.size = Pt(28)
run.font.color.rgb = BLUE
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BASE1 Restore Runbook')
run.font.size = Pt(22)
run.font.color.rgb = BLUE

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Version 1.0')
run.font.size = Pt(14)
run.font.color.rgb = GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Checkpoint: 29 May 2026 (Day 34)')
run.font.size = Pt(12)
run.font.color.rgb = GRAY

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Classification: Confidential — Ken Mun, CTO')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0xCF, 0x22, 0x2E)
run.bold = True

doc.add_paragraph()

today = datetime.date.today().strftime('%d %B %Y')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'Generated: {today}')
run.font.size = Pt(10)
run.font.color.rgb = GRAY

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (placeholder)
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled(doc, 'Table of Contents', 1)

toc_items = [
    '0. Prerequisites (Both Paths)',
    '1. Path A — Full Restore from Backup Files',
    '2. Path B — Full Restore from GitHub Repository',
    '3. Agent Configuration Reference',
    '4. Platform Configuration Reference',
    '5. Key Scripts Reference',
    '6. Cron Jobs Reference',
    '7. Verification Checklist',
    '8. Troubleshooting',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.runs[0].font.color.rgb = BLUE

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 0 — PREREQUISITES
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled(doc, '0. Prerequisites (Both Paths)', 1)

prereqs = [
    'macOS 26.x (Sequoia) on Apple Silicon (M4 or newer, minimum 24GB RAM)',
    'Stable internet connection',
    'Administrator (sudo) access',
    'Apple ID signed into App Store (required for Xcode Command Line Tools)',
    'Tailscale account credentials',
    'GitHub account: kenmun (personal) with Personal Access Token',
    'GitHub PAT must have scopes: repo, read:org, workflow',
    'Ollama Cloud Pro/Max account credentials',
    'Backup media (NAS path or external drive) — for Path A only',
]
for item in prereqs:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 1 — PATH A: FULL RESTORE FROM BACKUP
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled(doc, '1. Path A — Full Restore from Backup Files', 1)
p = doc.add_paragraph('This path restores the platform from a complete backup (NAS or external drive). Use this when you have a full file-level backup of ~/.openclaw/.')
p.runs[0].font.color.rgb = GRAY
p.runs[0].font.italic = True

# A1
add_heading_styled(doc, 'Step A1: System Preparation', 2)
doc.add_paragraph('Install Xcode Command Line Tools:', style='List Bullet')
add_code_block(doc, 'xcode-select --install')
doc.add_paragraph('Install Homebrew:', style='List Bullet')
add_code_block(doc, '/bin/bash -c "$(curl -fsSL https://raw.githubusercontent.com/Homebrew/install/HEAD/install.sh)"')

# A2
add_heading_styled(doc, 'Step A2: Install Core Dependencies', 2)
add_code_block(doc, '''brew install node@25 colima docker docker-compose jq gh tailscale rsync python@3.14
brew install --cask ollama''')

# A3
add_heading_styled(doc, 'Step A3: Install OpenClaw', 2)
add_code_block(doc, 'npm install -g openclaw@2026.5.12')

# A4
add_heading_styled(doc, 'Step A4: Install Python Platform Packages', 2)
add_code_block(doc, 'pip3 install --break-system-packages python-docx openpyxl fpdf pdfplumber pypdf python-pptx pandas requests GitPython opentimestamps-client')

# A5
add_heading_styled(doc, 'Step A5: Restore Configuration from Backup', 2)
doc.add_paragraph('Copy the following from your backup media (NAS/external drive) to the new Mac:')
items = [
    ('~/.openclaw/openclaw.json', 'Platform configuration — agents, models, channels, API keys'),
    ('~/.openclaw/workspace/', 'Full workspace (~270MB) — scripts, state, docs, memory, rules'),
    ('~/.openclaw/agents/', 'Agent runtime directory — session logs, tool grants'),
]
for path, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{path}')
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.bold = True
    p.add_run(f'  → {desc}')

doc.add_paragraph('Then fix permissions:')
add_code_block(doc, 'chmod -R 700 ~/.openclaw/agents/')

# A6
add_heading_styled(doc, 'Step A6: Restore GitHub Remote', 2)
add_code_block(doc, '''cd ~/.openclaw/workspace
git remote add origin https://github.com/kenmun/aevlith-nexus-platform.git
gh auth setup-git''')

# A7
add_heading_styled(doc, 'Step A7: Start Services', 2)
add_code_block(doc, '''# Start Colima (Docker runtime)
brew services start colima

# Start Tailscale
sudo tailscale up

# Start Ollama (local models)
ollama serve &

# Pull local models
ollama pull gemma4:26b     # 17GB — local inference model
ollama pull gemma4:e2b     # 7.2GB — benchmark model''')

# A8
add_heading_styled(doc, 'Step A8: Verify Installation', 2)
add_code_block(doc, '''# Check OpenClaw
openclaw status
openclaw gateway status

# Check agents (should show 14 agents)
openclaw agents list

# Check Docker
docker ps

# Check Tailscale
tailscale status

# Run health check (should pass 19/19)
bash ~/.openclaw/workspace/scripts/health-check.sh

# Run Warden drift check (should pass 9/9)
bash ~/.openclaw/workspace/scripts/model-drift-check.sh''')

# A9
add_heading_styled(doc, 'Step A9: Restart OpenClaw Gateway', 2)
add_code_block(doc, 'openclaw gateway restart')
doc.add_paragraph('The platform is now live. Proceed to the Verification Checklist (Section 7).')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 2 — PATH B: FULL RESTORE FROM GITHUB
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled(doc, '2. Path B — Full Restore from GitHub Repository', 1)
p = doc.add_paragraph('This path restores platform source code from the private GitHub repo. Requires manual restoration of openclaw.json (not in repo — gitignored for security).')
p.runs[0].font.color.rgb = GRAY
p.runs[0].font.italic = True

# B1-B4
add_heading_styled(doc, 'Steps B1–B4: System Preparation & Dependencies', 2)
doc.add_paragraph('Identical to Steps A1 through A4 above.')
doc.add_paragraph('Complete Steps A1, A2, A3, and A4 from Path A before proceeding.')

# B5
add_heading_styled(doc, 'Step B5: Clone the Platform Repository', 2)
add_code_block(doc, '''# Authenticate with GitHub (kenmun account)
gh auth login --hostname github.com

# Clone the private repo
mkdir -p ~/.openclaw
cd ~/.openclaw
git clone https://github.com/kenmun/aevlith-nexus-platform.git workspace''')

# B6
add_heading_styled(doc, 'Step B6: Configure OpenClaw', 2)
doc.add_paragraph('⚠️ CRITICAL: openclaw.json is NOT in the repository (gitignored for security).')
doc.add_paragraph('You have two options:')
doc.add_paragraph('Option 1 — Restore from backup file:', style='List Bullet')
p = doc.add_paragraph()
run = p.add_run('Copy openclaw.json from your backup to ~/.openclaw/openclaw.json')
run.font.name = 'Courier New'
run.font.size = Pt(9)
doc.add_paragraph('Option 2 — Manual reconfigure:', style='List Bullet')
doc.add_paragraph('Run openclaw init and manually configure 14 agents using the Agent Configuration Reference in Section 3.')

# B7
add_heading_styled(doc, 'Step B7: Recreate Agent Directory Structure', 2)
add_code_block(doc, '''mkdir -p ~/.openclaw/agents/{main,business,security,legal,qa,governance,infra,architect,platform-arch,biz-process,change-mgt,ahsoka,luthen,social}
mkdir -p ~/.openclaw/canvas/documents
mkdir -p ~/.openclaw/completions
mkdir -p ~/.openclaw/credentials''')

# B8-B9
add_heading_styled(doc, 'Steps B8–B9: GitHub Remote & Services', 2)
doc.add_paragraph('Identical to Steps A6 through A9 above (GitHub remote, services, verify, restart).')
doc.add_paragraph('Complete A6, A7, A8, and A9 from Path A to finish the restore.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 3 — AGENT CONFIGURATION REFERENCE
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled(doc, '3. Agent Configuration Reference', 1)
p = doc.add_paragraph('Use this reference if manually reconfiguring agents. All 14 agents with their role, model, and tier assignment.')
p.runs[0].font.color.rgb = GRAY
p.runs[0].font.italic = True

agents = [
    ('main', 'Yoda 🟢 — Lead Orchestrator', 'ollama/deepseek-v4-pro:cloud', 'userFacing'),
    ('business', 'Aria 🔵 — Business Lead', 'ollama/deepseek-v4-pro:cloud', 'userFacing'),
    ('security', 'Shield 🛡️ — Security Governance', 'ollama/gemma4:31b-cloud', 'backend'),
    ('legal', 'Lex ⚖️ — Legal Governance', 'ollama/gemma4:31b-cloud', 'backend'),
    ('qa', 'Sage 🧪 — QA Governance', 'ollama/gemma4:31b-cloud', 'backend'),
    ('governance', 'Warden 🔍 — Model Compliance', 'ollama/gemma4:31b-cloud', 'backend'),
    ('infra', 'Forge 🏗️ — Infra/Ops', 'ollama/gemma4:31b-cloud', 'backend'),
    ('architect', 'Atlas 🏛️ — Enterprise Architect', 'ollama/gemma4:31b-cloud', 'backend'),
    ('platform-arch', 'Thrawn — AI Platform Architect', 'ollama/gemma4:31b-cloud', 'backend'),
    ('biz-process', 'Lando 🟡 — Business Process', 'ollama/gemma4:31b-cloud', 'backend'),
    ('change-mgt', 'Mon Mothma 🌟 — Change Management', 'ollama/gemma4:31b-cloud', 'backend'),
    ('ahsoka', 'Ahsoka 🤍 — AI Consultant', 'ollama/gemma4:31b-cloud', 'backend'),
    ('luthen', 'Luthen 🔍 — Marketing Intelligence', 'ollama/gemma4:31b-cloud', 'backend'),
    ('social', 'Spark ✨ — Social & Content', 'ollama/kimi-k2.6:cloud', 'backend'),
]

table = doc.add_table(rows=len(agents)+1, cols=4)
table.style = 'Table Grid'
# Header
for i, h in enumerate(['Agent ID', 'Role', 'Primary Model', 'Tier']):
    cell = table.rows[0].cells[i]
    cell.text = h
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size = Pt(9)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(cell, '0969DA')

for i, (agent_id, role, model, tier) in enumerate(agents):
    row = table.rows[i+1]
    row.cells[0].text = agent_id
    row.cells[1].text = role
    row.cells[2].text = model
    row.cells[3].text = tier
    for j in range(4):
        for para in row.cells[j].paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
    if i % 2 == 0:
        for j in range(4):
            set_cell_shading(row.cells[j], 'F6F8FA')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 4 — PLATFORM CONFIGURATION
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled(doc, '4. Platform Configuration Reference', 1)

versions = [
    ('OpenClaw', '2026.5.12'),
    ('Node.js', 'v25.9.0'),
    ('Python', '3.14.5'),
    ('macOS', '26.4.1 (Sequoia)'),
    ('Homebrew', '5.1.14'),
    ('Docker / Colima', 'Colima 0.10.1'),
    ('Ollama', '0.23.2'),
    ('Tailscale', '1.96.5'),
    ('Git', '2.50.1 (Apple Git-155)'),
]

table = doc.add_table(rows=len(versions)+1, cols=2)
table.style = 'Table Grid'
for i, h in enumerate(['Component', 'Version']):
    cell = table.rows[0].cells[i]
    cell.text = h
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size = Pt(10)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(cell, '0969DA')

for i, (comp, ver) in enumerate(versions):
    row = table.rows[i+1]
    row.cells[0].text = comp
    row.cells[1].text = ver
    for j in range(2):
        for para in row.cells[j].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
    if i % 2 == 0:
        for j in range(2):
            set_cell_shading(row.cells[j], 'F6F8FA')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 5 — KEY SCRIPTS
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled(doc, '5. Key Scripts Reference', 1)

scripts_list = [
    ('scripts/health-check.sh', 'System health (19 checks, runs every 5 min)'),
    ('scripts/model-drift-check.sh', 'Warden compliance — agent model vs policy'),
    ('scripts/cost-tracker.sh', 'Daily cost calculation from session logs'),
    ('scripts/auto-heal.sh', 'Overnight auto-repair (01:00 MYT, 19 checks)'),
    ('scripts/run-diagnostics.sh', 'Full diagnostics (7 phases)'),
    ('scripts/ticket.sh', 'Ticket lifecycle management + Notion sync'),
    ('scripts/changelog-append.sh', 'CHG change control entries'),
    ('scripts/journal-append.sh', 'Daily journal inline writes'),
    ('scripts/generate-mission-control.sh', 'Dashboard data generation'),
    ('scripts/warden-cron.sh', 'Warden compliance cron runner'),
    ('scripts/gemma4-benchmark.sh', 'Model benchmark (gemma4:26b/e2b)'),
    ('scripts/cron-health-check.sh', 'Cron job health monitoring'),
    ('scripts/budget-check.sh', 'Daily budget spend vs cap'),
    ('scripts/owl-compliance-check.sh', 'OWL execution contract compliance'),
]

table = doc.add_table(rows=len(scripts_list)+1, cols=2)
table.style = 'Table Grid'
for i, h in enumerate(['Script', 'Purpose']):
    cell = table.rows[0].cells[i]
    cell.text = h
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size = Pt(9)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(cell, '0969DA')

for i, (script, purpose) in enumerate(scripts_list):
    row = table.rows[i+1]
    cell0 = row.cells[0]
    cell0.text = ''
    p = cell0.paragraphs[0]
    run = p.add_run(script)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    row.cells[1].text = purpose
    for j in range(2):
        for para in row.cells[j].paragraphs:
            for run in para.runs:
                if run.font.name != 'Courier New':
                    run.font.size = Pt(9)
    if i % 2 == 0:
        for j in range(2):
            set_cell_shading(row.cells[j], 'F6F8FA')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 6 — CRON JOBS
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled(doc, '6. Cron Jobs Reference', 1)
p = doc.add_paragraph('Key scheduled jobs to verify are running after restore. Managed via OpenClaw Gateway cron system.')
p.runs[0].font.color.rgb = GRAY
p.runs[0].font.italic = True

crons = [
    ('Warden Model Compliance Check', 'Hourly at :07 past', '83accf7b'),
    ('Gateway Health Check (silent)', 'Every 5 minutes', 'c65ace85'),
    ('Morning Stand-Up', 'Daily 08:00 MYT', '3c279099'),
    ('Auto-Heal', 'Daily 01:00 MYT', '—'),
    ('End-of-Day Journal', 'Daily 23:55 MYT', '4d926b2c'),
    ('Blog Generator', 'Daily 00:05 MYT', 'a027fd60'),
    ('Drive Sync', 'Daily 00:30 MYT', 'c5a3911d'),
    ('Task Monitor', 'Every 5 minutes', '637ecb12'),
    ('Task Queue Processor', 'Every 5 minutes', 'a89d00ef'),
    ('Aria→Ken Relay Queue', 'Every 5 minutes', '7a28cc83'),
    ('Mission Control Refresh', 'Every 15 minutes', 'd32f2b9a'),
    ('Observability Collector', 'Every 5 minutes', 'd3b1e203'),
    ('TZ Drift Monitor', 'Every 30 minutes', '9ce7f295'),
    ('PG Sync Check', 'Hourly', 'f7668f6a'),
]

table = doc.add_table(rows=len(crons)+1, cols=3)
table.style = 'Table Grid'
for i, h in enumerate(['Cron Job', 'Schedule', 'ID']):
    cell = table.rows[0].cells[i]
    cell.text = h
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size = Pt(9)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(cell, '0969DA')

for i, (name, schedule, cid) in enumerate(crons):
    row = table.rows[i+1]
    row.cells[0].text = name
    row.cells[1].text = schedule
    row.cells[2].text = cid
    for j in range(3):
        for para in row.cells[j].paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
    if i % 2 == 0:
        for j in range(3):
            set_cell_shading(row.cells[j], 'F6F8FA')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 7 — VERIFICATION CHECKLIST
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled(doc, '7. Verification Checklist', 1)
p = doc.add_paragraph('After completing either restore path, verify ALL items below before declaring the platform operational.')
p.runs[0].font.color.rgb = GRAY
p.runs[0].font.italic = True

checks = [
    'openclaw status shows gateway running',
    'openclaw gateway status returns OK',
    'All 14 agents listed in openclaw agents list',
    'bash scripts/health-check.sh passes all 19 checks',
    'bash scripts/model-drift-check.sh returns 9/9 PASS',
    'Tailscale connected: tailscale status shows active',
    'Docker running: docker ps returns without error',
    'Colima active: colima status shows running',
    'GitHub remote: git remote -v shows kenmun/aevlith-nexus-platform',
    'Ollama local models: ollama list shows gemma4:26b, gemma4:e2b',
    'Ollama Cloud accessible: ollama run gemma4:31b-cloud "hello" responds',
    'Git push works: make test commit and git push origin main succeeds',
    'Journal inline writes: bash scripts/journal-append.sh succeeds',
    'CHG logging: bash scripts/changelog-append.sh succeeds',
    'Cost tracker: bash scripts/cost-tracker.sh runs without error',
]

table = doc.add_table(rows=len(checks)+1, cols=2)
table.style = 'Table Grid'
for i, h in enumerate(['✓', 'Check']):
    cell = table.rows[0].cells[i]
    cell.text = h
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size = Pt(10)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(cell, '0969DA')

for i, check in enumerate(checks):
    row = table.rows[i+1]
    row.cells[0].text = '☐'
    row.cells[1].text = check
    row.cells[0].width = Cm(1.0)
    for j in range(2):
        for para in row.cells[j].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
    if i % 2 == 0:
        for j in range(2):
            set_cell_shading(row.cells[j], 'F6F8FA')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 8 — TROUBLESHOOTING
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled(doc, '8. Troubleshooting', 1)

issues = [
    ('"ollama not found" or "command not found: ollama"',
     'brew install --cask ollama\nollama serve &'),
    ('"gh auth fails" or "authentication failed"',
     'Ensure GitHub PAT has scopes: repo, read:org, workflow\nGenerate at: https://github.com/settings/tokens'),
    ('"Docker daemon not running"',
     'colima start\ndocker ps  # verify'),
    ('"Tailscale not connected"',
     'sudo tailscale up\ntailscale status  # verify'),
    ('"openclaw.json missing"',
     'Restore from backup to ~/.openclaw/openclaw.json OR reconfigure manually via openclaw init'),
    ('"gemma4:26b fails" or OOM error',
     'Model needs ~17GB RAM. On 24GB Mac Mini: close other applications first.\nIf still failing: use gemma4:e2b (7.2GB) or skip local models.'),
    ('"Port 11434 already in use"',
     'pkill ollama\nsleep 2\nollama serve &'),
    ('"Permission denied" on agent directories',
     'chmod -R 700 ~/.openclaw/agents/'),
    ('Git push fails with authentication error',
     'gh auth setup-git\ngit config --global credential.helper osxkeychain'),
    ('openclaw gateway not starting',
     'Check logs: openclaw gateway logs\nVerify Node.js version: node --version (must be v25.x)\nReinstall: npm install -g openclaw@2026.5.12'),
]

for i, (problem, solution) in enumerate(issues):
    add_heading_styled(doc, f'Issue {i+1}: {problem}', 3)
    doc.add_paragraph('Solution:')
    add_code_block(doc, solution)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════
# FOOTER
# ═══════════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— End of Document —')
run.font.size = Pt(12)
run.font.color.rgb = GRAY
run.italic = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BASE1 Checkpoint — 29 May 2026 | Aevlith Nexus Platform | Ken Mun, CTO | Confidential')
run.font.size = Pt(9)
run.font.color.rgb = GRAY

# ── Headers & Footers ────────────────────────────────────────────────────────
for section in doc.sections:
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.text = 'Aevlith Nexus — BASE1 Restore Runbook v1.0'
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.text = 'Confidential — Ken Mun, CTO'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY

# ── Save ─────────────────────────────────────────────────────────────────────
output = '/Users/ainchorsoc2a/.openclaw/workspace/docs/deliverables/Aevlith-Nexus-BASE1-Restore-Runbook-v1.0.docx'
doc.save(output)
print(f'DOCX saved: {output}')

import os
size_mb = os.path.getsize(output) / (1024*1024)
print(f'Size: {size_mb:.1f} MB')
