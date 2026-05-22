#!/usr/bin/env python3
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx_template(path, title, subtitle, author):
    doc = Document()
    
    # Cover Page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.bold = True
    run.font.size = Pt(24)
    
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para.add_run(subtitle).font.size = Pt(14)
    
    doc.add_paragraph("\n" * 5)
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_para.add_run(f"Author: {author}\nDate: 2026-05-21\n[AInCHORS PLACEHOLDER LOGO]")

    doc.add_page_break()
    
    # TOC Placeholder
    doc.add_heading('Table of Contents', level=1)
    doc.add_paragraph("[TOC generated automatically by Word]")
    
    # Content
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph('This is a placeholder for the introduction section. This document uses the standard AInCHORS layout.')
    
    doc.add_heading('Project Scope', level=2)
    doc.add_paragraph('Details about the project scope go here.')
    
    doc.add_heading('Implementation Plan', level=2)
    doc.add_paragraph('Steps for implementation are listed here.')
    
    # Callout Box (Simulation via paragraph styling)
    doc.add_paragraph("\n")
    callout = doc.add_paragraph()
    callout.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = callout.add_run('IMPORTANT: This is a styled callout box for key highlights.')
    run.italic = True
    
    doc.save(path)

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 5:
        print("Usage: python3 generate_docx.py <path> <title> <subtitle> <author>")
        sys.exit(1)
    create_docx_template(sys.argv[1], sys.argv[2], sys.argv[3], sys.argv[4])
