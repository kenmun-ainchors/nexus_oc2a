#!/usr/bin/env python3
"""
proposal.py — AInchors Consulting Proposal (DOCX)
Usage: python3 proposal.py --title "Title" --output /path/out.docx [--data /path/data.json]
"""
import argparse
import json
import sys
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def add_header_para(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(0x0D, 0x11, 0x17)
    return p


def add_footer(doc, prepared_by):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"AInchors — ainchors.com — Prepared by {prepared_by}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--title', required=True)
    parser.add_argument('--output', required=True)
    parser.add_argument('--data', default=None)
    args = parser.parse_args()

    # Load data
    data = {}
    if args.data:
        with open(args.data) as f:
            data = json.load(f)

    client_name = data.get('clientName', 'Valued Client')
    doc_date = data.get('date', date.today().strftime('%d %B %Y'))
    scope_items = data.get('scopeItems', [
        'Discovery & stakeholder interviews',
        'Current-state process mapping',
        'AI readiness assessment',
        'Recommendations report',
    ])
    investment_rows = data.get('investmentRows', [
        {'item': 'Discovery Workshop', 'qty': 1, 'rate': 5000, 'total': 5000},
        {'item': 'Process Analysis', 'qty': 3, 'rate': 2500, 'total': 7500},
        {'item': 'Final Report & Presentation', 'qty': 1, 'rate': 3500, 'total': 3500},
    ])
    prepared_by = data.get('preparedBy', 'Ahsoka — AI Transformation Consultant')

    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── COVER ──────────────────────────────────────────────────────────────────
    # Logo header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('AInchors')
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0D, 0x11, 0x17)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('AI Transformation Consulting')
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    doc.add_paragraph()

    # Title
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(args.title)
    run3.font.size = Pt(24)
    run3.font.bold = True

    doc.add_paragraph()

    # Client / date / confidential
    for line in [f'Prepared for: {client_name}', f'Date: {doc_date}', 'CONFIDENTIAL']:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(12)
        if line == 'CONFIDENTIAL':
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_page_break()

    # ── EXECUTIVE SUMMARY ─────────────────────────────────────────────────────
    add_header_para(doc, '1. Executive Summary', level=1)
    doc.add_paragraph(
        'AInchors has been engaged to assess and accelerate your organisation\'s AI transformation journey. '
        'This proposal outlines our recommended approach, scope of work, and investment required to deliver '
        'measurable outcomes aligned to your strategic priorities. '
        'Our team brings deep expertise in AI strategy, process automation, and change management.'
    )

    # ── SCOPE OF WORK ─────────────────────────────────────────────────────────
    add_header_para(doc, '2. Scope of Work', level=1)
    for i, item in enumerate(scope_items, 1):
        doc.add_paragraph(f'{i}. {item}', style='List Number')

    # ── APPROACH & METHODOLOGY ────────────────────────────────────────────────
    add_header_para(doc, '3. Approach & Methodology', level=1)
    phases = [
        ('Phase 1 — Discover', 'Stakeholder interviews, data collection, and current-state process mapping to establish a baseline.'),
        ('Phase 2 — Design', 'Co-design of AI-enabled future-state processes, technology selection, and roadmap development.'),
        ('Phase 3 — Deliver', 'Implementation support, change management, training, and post-go-live review.'),
    ]
    for phase_title, phase_desc in phases:
        p = doc.add_paragraph()
        run = p.add_run(phase_title)
        run.font.bold = True
        doc.add_paragraph(phase_desc)

    # ── INVESTMENT ────────────────────────────────────────────────────────────
    add_header_para(doc, '4. Investment', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for i, hdr in enumerate(['Item', 'Qty', 'Rate (AUD)', 'Total (AUD)']):
        hdr_cells[i].text = hdr
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr_cells[i], '4472C4')

    total_sum = 0
    for row in investment_rows:
        cells = table.add_row().cells
        cells[0].text = str(row.get('item', ''))
        cells[1].text = str(row.get('qty', ''))
        cells[2].text = f"${row.get('rate', 0):,}"
        cells[3].text = f"${row.get('total', 0):,}"
        total_sum += int(row.get('total', 0))

    # Total row
    tot_cells = table.add_row().cells
    tot_cells[0].text = 'Total'
    tot_cells[0].paragraphs[0].runs[0].font.bold = True
    tot_cells[3].text = f'${total_sum:,}'
    tot_cells[3].paragraphs[0].runs[0].font.bold = True

    # ── TERMS & NEXT STEPS ────────────────────────────────────────────────────
    add_header_para(doc, '5. Terms & Next Steps', level=1)
    doc.add_paragraph(
        'Payment terms: 50% upon engagement, 50% on delivery of final report. '
        'To proceed, please sign and return this proposal within 14 days. '
        'A kick-off meeting will be scheduled within 5 business days of acceptance.'
    )
    steps = [
        'Sign and return this proposal',
        'Complete initial payment (50%)',
        'Kick-off meeting scheduled',
        'Discovery phase commences',
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Bullet')

    # ── FOOTER ────────────────────────────────────────────────────────────────
    add_footer(doc, prepared_by)

    doc.save(args.output)
    print(f'Proposal saved: {args.output}')


if __name__ == '__main__':
    main()
